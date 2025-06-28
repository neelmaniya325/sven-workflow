from fastapi import FastAPI, UploadFile, HTTPException
from fastapi.responses import FileResponse
import mammoth
import re
import os
import tempfile
from pathlib import Path
import zipfile
from docx import Document


app = FastAPI()


def anonymize_text(text: str) -> str:
    """
    Anonymize text by replacing PII with placeholder tags.
    """
    # Email addresses (must come before general patterns)
    text = re.sub(r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b', '[Email address]', text)
    
    # Phone numbers (various formats)
    phone_patterns = [
        r'\b\+?[\d\s\-\(\)]{10,}\b',  # General phone pattern
        r'\b\d{3}[-.\s]?\d{3}[-.\s]?\d{4}\b',  # US format
        r'\b\(\d{3}\)\s?\d{3}[-.\s]?\d{4}\b',  # (123) 456-7890
        r'\b\d{2,4}[-.\s]?\d{2,4}[-.\s]?\d{2,4}[-.\s]?\d{2,4}\b'  # International
    ]
    for pattern in phone_patterns:
        text = re.sub(pattern, '[Telephone number]', text)
    
    # Social Security Numbers
    text = re.sub(r'\b\d{3}-\d{2}-\d{4}\b', '[Social security number]', text)
    text = re.sub(r'\b\d{9}\b', '[Social security number]', text)
    
    # Bank account and similar long numbers
    text = re.sub(r'\b\d{10,20}\b', '[Bank account number]', text)
    
    # Insurance numbers
    text = re.sub(r'\b[A-Z]{2,3}\d{6,12}\b', '[Insurance number]', text)
    
    # License plate numbers
    text = re.sub(r'\b[A-Z]{1,3}[-\s]?\d{3,4}[-\s]?[A-Z]{0,3}\b', '[License plate number]', text)
    
    # IP addresses
    text = re.sub(r'\b(?:\d{1,3}\.){3}\d{1,3}\b', '[IP address]', text)
    
    # URLs and domains
    url_patterns = [
        r'https?://[^\s]+',
        r'www\.[^\s]+',
        r'\b[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}\b'
    ]
    for pattern in url_patterns:
        text = re.sub(pattern, '[URL or account]', text)
    
    # Dates (various formats)
    date_patterns = [
        r'\b\d{1,2}[./\-]\d{1,2}[./\-]\d{2,4}\b',  # MM/DD/YYYY, DD.MM.YYYY
        r'\b\d{2,4}[./\-]\d{1,2}[./\-]\d{1,2}\b',  # YYYY-MM-DD
        r'\b(?:January|February|March|April|May|June|July|August|September|October|November|December)\s+\d{1,2},?\s+\d{2,4}\b',
        r'\b\d{1,2}\s+(?:January|February|March|April|May|June|July|August|September|October|November|December)\s+\d{2,4}\b'
    ]
    for pattern in date_patterns:
        text = re.sub(pattern, '[Date of birth]', text, flags=re.IGNORECASE)
    
    # ZIP/Postal codes
    text = re.sub(r'\b\d{5}(-\d{4})?\b', '[Postal code]', text)  # US ZIP
    text = re.sub(r'\b[A-Z]\d[A-Z]\s?\d[A-Z]\d\b', '[Postal code]', text)  # Canadian
    text = re.sub(r'\b\d{4,6}\b', '[Postal code]', text)  # General postal codes
    
    # Street addresses and house numbers
    address_patterns = [
        r'\b\d+\s+[A-Za-z\s]+(Street|St|Avenue|Ave|Road|Rd|Drive|Dr|Lane|Ln|Boulevard|Blvd|Circle|Cir|Court|Ct|Place|Pl)\b',
        r'\b\d+\s+[A-Za-z\s]+\s+\d+\b'  # House number + street name + apartment
    ]
    for pattern in address_patterns:
        text = re.sub(pattern, '[Street and house number]', text, flags=re.IGNORECASE)
    
    # Case/File numbers
    text = re.sub(r'\b(?:Case|File|Ref|Reference)[\s#:]*[A-Z0-9\-]+\b', '[Case number]', text, flags=re.IGNORECASE)
    
    # First name patterns (common first names)
    first_names = r'\b(?:John|Jane|Michael|Sarah|David|Lisa|Robert|Mary|James|Patricia|William|Jennifer|Richard|Linda|Charles|Elizabeth|Thomas|Barbara|Christopher|Susan|Daniel|Jessica|Matthew|Nancy|Anthony|Dorothy|Mark|Karen|Donald|Helen|Steven|Michelle|Paul|Sandra|Andrew|Donna|Joshua|Carol|Kenneth|Ruth|Kevin|Sharon|Brian|Michelle|George|Laura|Edward|Sarah|Ronald|Kimberly|Timothy|Deborah|Jason|Dorothy|Jeffrey|Lisa|Ryan|Nancy|Jacob|Karen|Gary|Betty|Nicholas|Helen|Eric|Sandra|Jonathan|Donna|Stephen|Carol|Larry|Ruth|Justin|Sharon|Scott|Michelle|Brandon|Laura|Benjamin|Sarah|Samuel|Kimberly|Gregory|Deborah|Alexander|Dorothy|Patrick|Lisa|Frank|Nancy|Raymond|Karen|Jack|Betty|Dennis|Helen|Jerry|Sandra|Tyler|Donna|Aaron|Carol|Jose|Ruth|Henry|Sharon|Adam|Michelle|Douglas|Laura|Nathan|Sarah|Peter|Kimberly|Zachary|Deborah|Kyle|Dorothy|Noah|Lisa)\b'
    text = re.sub(first_names, '[First name]', text, flags=re.IGNORECASE)
    
    # Professional titles with names
    titles = r'\b(?:Dr|Doctor|Prof|Professor|Mr|Mrs|Ms|Miss|Sir|Madam|Attorney|Lawyer|Judge|Officer|Detective|Sergeant|Lieutenant|Captain|Major|Colonel|General|President|Director|Manager|CEO|CFO|CTO|Principal|Dean|Chancellor)\.?\s+[A-Z][a-z]+(?:\s+[A-Z][a-z]+)?\b'
    text = re.sub(titles, '[Professional\'s name]', text, flags=re.IGNORECASE)
    
    # Institution names
    institution_patterns = [
        r'\b[A-Z][a-z]+\s+(?:University|College|School|Hospital|Medical Center|Clinic|Corporation|Company|Inc|LLC|Ltd)\b',
        r'\b(?:University|College|School|Hospital|Medical Center|Clinic)\s+of\s+[A-Z][a-z]+\b'
    ]
    for pattern in institution_patterns:
        text = re.sub(pattern, '[Institution name]', text, flags=re.IGNORECASE)
    
    # Job titles
    job_titles = r'\b(?:Manager|Director|Supervisor|Coordinator|Specialist|Analyst|Engineer|Developer|Designer|Consultant|Administrator|Assistant|Secretary|Clerk|Technician|Operator|Worker|Employee|Staff|Representative|Agent|Officer|Executive|President|Vice President|CEO|CFO|CTO|Principal|Teacher|Professor|Doctor|Nurse|Therapist|Counselor|Social Worker|Case Worker)\b'
    text = re.sub(job_titles, '[Occupational title]', text, flags=re.IGNORECASE)
    
    # Court names
    text = re.sub(r'\b[A-Z][a-z]+\s+(?:Court|Courthouse|Tribunal)\b', '[Court name]', text)
    
    # GPS coordinates
    text = re.sub(r'\b\-?\d{1,3}\.\d+,\s*\-?\d{1,3}\.\d+\b', '[Geo data]', text)
    
    # Patient/Client numbers
    text = re.sub(r'\b(?:Patient|Client|ID)[\s#:]*\d+\b', '[Patient number]', text, flags=re.IGNORECASE)
    
    # Last names (common surnames) - be careful with this as it might over-match
    last_names = r'\b(?:Smith|Johnson|Williams|Brown|Jones|Garcia|Miller|Davis|Rodriguez|Martinez|Hernandez|Lopez|Gonzalez|Wilson|Anderson|Thomas|Taylor|Moore|Jackson|Martin|Lee|Perez|Thompson|White|Harris|Sanchez|Clark|Ramirez|Lewis|Robinson|Walker|Young|Allen|King|Wright|Scott|Torres|Nguyen|Hill|Flores|Green|Adams|Nelson|Baker|Hall|Rivera|Campbell|Mitchell|Carter|Roberts)\b'
    text = re.sub(last_names, '[Last name]', text, flags=re.IGNORECASE)
    
    return text


def anonymize_docx_file(input_path: str, output_path: str) -> None:
    """
    Anonymize a DOCX file while preserving formatting.
    """
    doc = Document(input_path)
    
    # Anonymize paragraphs
    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            anonymized_text = anonymize_text(paragraph.text)
            paragraph.text = anonymized_text
    
    # Anonymize tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    cell.text = anonymize_text(cell.text)
    
    # Save the anonymized document
    doc.save(output_path)


def create_anonymized_file(file_path: str, file_extension: str) -> str:
    """
    Create anonymized version of the file based on its type.
    """
    base_name = os.path.splitext(os.path.basename(file_path))[0]
    output_path = f"{base_name}_anonymized{file_extension}"
    
    if file_extension.lower() == '.docx':
        anonymize_docx_file(file_path, output_path)
    else:
        # For other file types, extract text and create a new document
        with open(file_path, "rb") as f:
            result = mammoth.extract_raw_text(f)
        
        anonymized_content = anonymize_text(result.value)
        
        # Create a new DOCX file with anonymized content
        doc = Document()
        
        # Split content into paragraphs and add to document
        paragraphs = anonymized_content.split('\n')
        for para_text in paragraphs:
            if para_text.strip():
                doc.add_paragraph(para_text)
        
        output_path = f"{base_name}_anonymized.docx"
        doc.save(output_path)
    
    return output_path


@app.post('/upload-file')
async def upload_and_anonymize_file(file: UploadFile):
    """
    Upload a file, anonymize it, and return the anonymized file for download.
    """
    try:
        # Validate file type
        if not file.filename:
            raise HTTPException(status_code=400, detail="No file provided")
        
        file_extension = os.path.splitext(file.filename)[1].lower()
        supported_extensions = ['.docx', '.doc']
        
        if file_extension not in supported_extensions:
            raise HTTPException(
                status_code=400, 
                detail=f"Unsupported file type. Supported types: {', '.join(supported_extensions)}"
            )
        
        # Create temporary directory
        temp_dir = tempfile.mkdtemp()
        
        try:
            # Save uploaded file
            input_file_path = os.path.join(temp_dir, file.filename)
            content = await file.read()
            
            with open(input_file_path, "wb") as f:
                f.write(content)
            
            # Create anonymized file
            anonymized_file_path = create_anonymized_file(input_file_path, file_extension)
            
            # Move anonymized file to temp directory for cleanup
            final_anonymized_path = os.path.join(temp_dir, os.path.basename(anonymized_file_path))
            if anonymized_file_path != final_anonymized_path:
                os.rename(anonymized_file_path, final_anonymized_path)
            
            # Return the anonymized file for download
            return FileResponse(
                path=final_anonymized_path,
                filename=os.path.basename(final_anonymized_path),
                media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                headers={"Content-Disposition": f"attachment; filename={os.path.basename(final_anonymized_path)}"}
            )
            
        except Exception as e:
            # Clean up temp directory in case of error
            import shutil
            shutil.rmtree(temp_dir, ignore_errors=True)
            raise e
            
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error processing file: {str(e)}")

